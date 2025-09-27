import base64
import os
import io
import logging
from typing import List, Dict, Any
from docx import Document
import openpyxl

logger = logging.getLogger(__name__)


class FileParser:
    """Minimal file parser for Vercel deployment - supports only essential formats"""
    
    @staticmethod
    def decode_base64_file(content_base64: str, filename: str, upload_dir: str) -> str:
        """Decode base64 content and save to file"""
        try:
            # Ensure upload directory exists
            os.makedirs(upload_dir, exist_ok=True)
            
            # Decode base64 content
            file_content = base64.b64decode(content_base64)
            
            # Create unique filename to avoid conflicts
            import uuid
            unique_filename = f"{uuid.uuid4()}_{filename}"
            file_path = os.path.join(upload_dir, unique_filename)
            
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"Successfully saved file: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error decoding base64 file {filename}: {str(e)}")
            raise
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            text_content = []
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            combined_text = "\n".join(text_content)
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return combined_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def parse_xlsx(file_path: str) -> str:
        """Extract data from XLSX file"""
        try:
            text_content = []
            workbook = openpyxl.load_workbook(file_path)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data if cell):
                        text_content.append(" | ".join(row_data))
            
            combined_text = "\n".join(text_content)
            logger.info(f"Successfully extracted data from XLSX: {file_path}")
            return combined_text
            
        except Exception as e:
            logger.error(f"Error parsing XLSX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get file extension in lowercase"""
        return os.path.splitext(filename)[1].lower()
    
    @staticmethod
    def parse_file(content_base64: str, filename: str, upload_dir: str) -> str:
        """Parse file based on extension and return extracted text"""
        try:
            # Decode and save file
            file_path = FileParser.decode_base64_file(content_base64, filename, upload_dir)
            
            # Get file extension
            extension = FileParser.get_file_extension(filename)
            
            # Parse based on file type (only supported formats)
            if extension == '.docx':
                return FileParser.parse_docx(file_path)
            elif extension in ['.xlsx', '.xls']:
                return FileParser.parse_xlsx(file_path)
            else:
                logger.warning(f"Unsupported file type in minimal mode: {extension}")
                return f"[File type {extension} not supported in minimal deployment mode]"
                
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            raise
        finally:
            # Clean up temporary file
            try:
                if 'file_path' in locals() and os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info(f"Cleaned up temporary file: {file_path}")
            except Exception as cleanup_error:
                logger.warning(f"Error cleaning up file {file_path}: {str(cleanup_error)}")


def parse_attachments(attachments: List[Dict[str, Any]], upload_dir: str) -> str:
    """Parse multiple attachments and return combined text"""
    try:
        combined_text = []
        
        for attachment in attachments:
            filename = attachment.get('filename', 'unknown')
            content_base64 = attachment.get('contentBase64', '')
            
            if not content_base64:
                logger.warning(f"No content provided for attachment: {filename}")
                continue
            
            try:
                logger.info(f"Processing attachment: {filename}")
                text_content = FileParser.parse_file(content_base64, filename, upload_dir)
                
                if text_content:
                    combined_text.append(f"=== {filename} ===\n{text_content}\n")
                else:
                    combined_text.append(f"=== {filename} ===\n[No text content extracted]\n")
                    
            except Exception as e:
                logger.error(f"Error processing attachment {filename}: {str(e)}")
                combined_text.append(f"=== {filename} ===\n[Error processing file: {str(e)}]\n")
        
        result = "\n".join(combined_text)
        logger.info(f"Successfully processed {len(attachments)} attachments")
        return result
        
    except Exception as e:
        logger.error(f"Error processing attachments: {str(e)}")
        raise
