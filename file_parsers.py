import base64
import os
import io
import logging
from typing import List, Dict, Any
import pdfplumber
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

# Optional imports for file types we may not support in production
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - Word document parsing disabled")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logger.warning("pandas not available - Excel parsing disabled")

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("pytesseract/PIL not available - OCR parsing disabled")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("openpyxl not available - Excel parsing disabled")


class FileParser:
    """Base class for file parsing operations"""
    
    @staticmethod
    def decode_base64_file(content_base64: str, filename: str, upload_dir: str) -> str:
        """Decode base64 content and save to file"""
        try:
            # Ensure upload directory exists
            os.makedirs(upload_dir, exist_ok=True)
            
            # Decode base64 content
            file_content = base64.b64decode(content_base64)
            
            # Create unique filename to avoid conflicts
            import uuid
            unique_filename = f"{uuid.uuid4()}_{filename}"
            file_path = os.path.join(upload_dir, unique_filename)
            
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"Successfully saved file: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error decoding base64 file {filename}: {str(e)}")
            raise
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF using pdfplumber"""
        try:
            text_content = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            combined_text = "\n".join(text_content)
            logger.info(f"Successfully extracted text from PDF: {file_path}")
            return combined_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            # Fallback to PyMuPDF if pdfplumber fails
            try:
                return FileParser.parse_pdf_pymupdf(file_path)
            except Exception as e2:
                logger.error(f"Error parsing PDF with PyMuPDF {file_path}: {str(e2)}")
                raise
    
    @staticmethod
    def parse_pdf_pymupdf(file_path: str) -> str:
        """Extract text from PDF using PyMuPDF as fallback"""
        try:
            text_content = []
            doc = fitz.open(file_path)
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text:
                    text_content.append(text)
            
            doc.close()
            combined_text = "\n".join(text_content)
            logger.info(f"Successfully extracted text from PDF using PyMuPDF: {file_path}")
            return combined_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF with PyMuPDF {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            return f"Word document parsing not available: {os.path.basename(file_path)} (python-docx not installed)"
        
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            combined_text = "\n".join(text_content)
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return combined_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def parse_xlsx(file_path: str) -> str:
        """Extract text from XLSX file"""
        if not PANDAS_AVAILABLE or not EXCEL_AVAILABLE:
            return f"Excel parsing not available: {os.path.basename(file_path)} (pandas/openpyxl not installed)"
        
        try:
            text_content = []
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Add sheet name as header
                text_content.append(f"Sheet: {sheet_name}")
                
                # Convert dataframe to text representation
                if not df.empty:
                    # Handle NaN values
                    df = df.fillna('')
                    
                    # Get column names
                    columns = list(df.columns)
                    text_content.append(" | ".join(str(col) for col in columns))
                    
                    # Get row data
                    for _, row in df.iterrows():
                        row_data = [str(cell) for cell in row.values]
                        text_content.append(" | ".join(row_data))
                
                text_content.append("")  # Add spacing between sheets
            
            combined_text = "\n".join(text_content)
            logger.info(f"Successfully extracted text from XLSX: {file_path}")
            return combined_text
            
        except Exception as e:
            logger.error(f"Error parsing XLSX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def parse_image(file_path: str) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            return f"Image OCR not available: {os.path.basename(file_path)} (pytesseract/PIL not installed)"
        
        try:
            # Open image
            image = Image.open(file_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            
            logger.info(f"Successfully extracted text from image using OCR: {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing image {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get file extension from filename"""
        return os.path.splitext(filename)[1].lower()
    
    @staticmethod
    def parse_file(content_base64: str, filename: str, upload_dir: str) -> str:
        """Parse file based on extension and return extracted text"""
        try:
            # Decode and save file
            file_path = FileParser.decode_base64_file(content_base64, filename, upload_dir)
            
            # Get file extension
            extension = FileParser.get_file_extension(filename)
            
            # Parse based on file type
            if extension == '.pdf':
                return FileParser.parse_pdf(file_path)
            elif extension == '.docx':
                return FileParser.parse_docx(file_path)
            elif extension in ['.xlsx', '.xls']:
                return FileParser.parse_xlsx(file_path)
            elif extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return FileParser.parse_image(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            raise
        finally:
            # Clean up temporary file
            try:
                if 'file_path' in locals() and os.path.exists(file_path):
                    os.remove(file_path)
                    logger.info(f"Cleaned up temporary file: {file_path}")
            except Exception as cleanup_error:
                logger.warning(f"Error cleaning up file {file_path}: {str(cleanup_error)}")


def parse_attachments(attachments: List[Dict[str, Any]], upload_dir: str) -> str:
    """Parse all attachments and return combined text"""
    all_text = []
    
    for attachment in attachments:
        try:
            filename = attachment.get('filename', '')
            content_base64 = attachment.get('contentBase64', '')
            
            if not filename or not content_base64:
                logger.warning(f"Skipping attachment with missing filename or content: {attachment}")
                continue
            
            logger.info(f"Parsing attachment: {filename}")
            extracted_text = FileParser.parse_file(content_base64, filename, upload_dir)
            
            if extracted_text.strip():
                all_text.append(f"=== {filename} ===")
                all_text.append(extracted_text)
                all_text.append("")  # Add spacing
            else:
                logger.info(f"No text extracted from: {filename}")
                
        except Exception as e:
            logger.error(f"Error processing attachment {attachment.get('filename', 'unknown')}: {str(e)}")
            continue
    
    return "\n".join(all_text)
